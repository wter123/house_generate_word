

from tkinter import *
import os
import docx,lxml
import xlwings
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import time
from multiprocessing import Pool
import psutil
from tkinter.ttk import Progressbar

import xlrd
import os
import docx,lxml
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import time
from tkinter.ttk import Progressbar
import xlrd
import re
from xlrd import open_workbook, xldate
from tkinter import scrolledtext  # 导入滚动文本框的模块


import xlrd
import os
import docx,lxml
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import time
from tkinter.ttk import Progressbar
import xlrd
import re
from xlrd import open_workbook, xldate
from tkinter import scrolledtext  # 导入滚动文本框的模块









from tkinter import scrolledtext  # 导入滚动文本框的模块
# start_time=time.time()
import xlwt
from xlutils import copy
# s1 = r"C:\\model\非危房"
from PIL import Image

class deal_word():
    def __init__(self):

        # 路径
        self._i=0
        self._s1 = ""
        self._docx_path1 = os.path.join(self._s1,r'模板\非危房模板.docx')
        self._docx_path2 = os.path.join(self._s1,r'模板\危房模板.docx')
        self._docx_folder = os.path.join(self._s1,"文件")
        self._excel_path=os.path.join(self._s1,r"模板\模板.xlsx")
        self._pictrue_folder = os.path.join(self._s1,"照片")
        self._excel_cell =  " "
        self._list_problem=[]
        self._ibc=0
        self._save_p=""
        self._serial=""

        # 变量
        # self.doc=""
        self._pictrue_count=0
    def _all_to_string(self,ko):
        if not ko:
            return ""
        elif type(ko) == int:
            str(ko)
        elif type(ko) == float:
            return str(int(ko))
        else:
            return str(ko)
    def _insert_word(self,excel_cell, i, doc):
        # 建筑物位置
        p = doc.tables[0].cell(0, 2).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 3).value))
        p.font.name = u"宋体"

        # 建筑物坐标
        c = "东经（度分秒）:" + excel_cell.cell(i, 1).value + "                  北纬（度分秒）:" + excel_cell.cell(i, 2).value
        p = doc.tables[0].cell(1, 2).paragraphs[0].clear().add_run(self._all_to_string(c))
        p.font.name = u"宋体"

        # 普查编号

        a = re.search("(.*?)[(（](.*?)[)）]", excel_cell.cell(i, 0).value)
        # print(a)
        if a:
            print(a.group(1), excel_cell.cell(i, 0).value)
            p = doc.tables[0].cell(2, 2).paragraphs[0].clear().add_run(
                self._all_to_string(a.group(1)))  # t
        else:
            p = doc.tables[0].cell(2, 2).paragraphs[0].clear().add_run(
                self._all_to_string(excel_cell.cell(i, 0).value))  # t
        p.font.name = u"宋体"

        # 建造年份
        # print(
        # xldate.xldate_as_datetime(excel_cell.cell(i, 4).value,0)
        # )
        if (str(excel_cell.cell(i, 4).value)):
            try:
                c = str(xldate.xldate_as_datetime(excel_cell.cell(i, 4).value, 0))
                cc = c.replace(" 00:00:00", "")
            except:
                cc =str(excel_cell.cell(i, 4).value)
        else:
            cc=""
        p = doc.tables[0].cell(2, 8).paragraphs[0].clear().add_run(self._all_to_string(cc))  # text="aaaa"
        p.font.name = u"宋体"

        def switch_test_item(item):
            switcher = {
                "住宅类": "A √住宅类   B 工业类   C 商业类   D 办公类   E 其他",
                "工业类": "A 住宅类   B√ 工业类   C 商业类   D 办公类   E 其他",
                "商业类": "A 住宅类   B 工业类   C√ 商业类   D 办公类   E 其他",
                "办公类": "A 住宅类   B 工业类   C 商业类   D √办公类   E 其他",
                "其他": "A 住宅类   B 工业类   C 商业类   D 办公类   E √其他",
            }
            return switcher.get(item.replace(" ", ""), "A 住宅类   B 工业类   C 商业类   D 办公类   E 其他")

        # 建筑物用途
        p = doc.tables[0].cell(3, 2).paragraphs[0].clear().add_run(
            switch_test_item(self._all_to_string(excel_cell.cell(i, 8).value)))  # text="aaaa"
        p.font.name = u"宋体"

        # 建筑层数
        p = doc.tables[0].cell(4, 2).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 5).value))  # text="aaaa"
        p.font.name = u"宋体"

        # 建筑面积
        p = doc.tables[0].cell(4, 6).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 6).value))  # text="aaaa"
        p.font.name = u"宋体"

        # 结构类型
        p = doc.tables[0].cell(4, 9).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 7).value))  # text="aaaa"
        p.font.name = u"宋体"

        # 安全责任人
        p = doc.tables[0].cell(5, 2).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 9).value))  # text="aaaa"
        p.font.name = u"宋体"

        # 联系电话
        p = doc.tables[0].cell(5, 6).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 10).value))  # text="aaaa"
        p.font.name = u"宋体"

        # 是否存在整体或局部倾斜、下沉情况
        p = doc.tables[0].cell(8, 7).paragraphs[0].clear().add_run(self._all_to_string(excel_cell.cell(i, 12).value))
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 11).value == u"是"):
            p = doc.tables[0].cell(8, 5).paragraphs[0].clear().add_run("是☑  否□")
            p.font.name = u"Segoe UI Symbol"
            pass
        elif (excel_cell.cell(i, 11).value == u"否"):
            p = doc.tables[0].cell(8, 5).paragraphs[0].clear().add_run("是□  否☑")
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(8, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充整体情况第一项")
            p.font.name = u"Segoe UI Symbol"

        # 周边是否存在地质灾害、洪水内涝等安全隐患
        p = doc.tables[0].cell(9, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 14).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 13).value == u"是"):
            p = doc.tables[0].cell(9, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
            pass
        elif (excel_cell.cell(i, 13).value == u"否"):
            p = doc.tables[0].cell(9, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(9, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充整体情况第二项")

            p.font.name = u"Segoe UI Symbol"

        # 周边是否存在地下管线、桩基、深基坑施工、爆破等可能被损坏的情况
        p = doc.tables[0].cell(10, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 16).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 15).value == u"是"):
            p = doc.tables[0].cell(10, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 15).value == u"否"):
            p = doc.tables[0].cell(10, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:

            p = doc.tables[0].cell(10, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充整体情况第三项")
            # [] if doc.tables[0].cell(10, 0).text in self._list_problem else self._list_problem.appendd(doc.tables[0].cell(10, 0).text)
            p.font.name = u"Segoe UI Symbol"

        # 是否存在下沉、裂缝、变形、渗水等
        p = doc.tables[0].cell(11, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 18).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 17).value == u"是"):
            p = doc.tables[0].cell(11, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 17).value == u"否"):
            p = doc.tables[0].cell(11, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(11, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充基础（地坪）第一项")
            p.font.name = u"Segoe UI Symbol"

        # 梁柱板、墙体及钢结构等构件是否存在破损、变形、裂缝、渗水、脱落、腐蚀等情况
        p = doc.tables[0].cell(12, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 20).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 19).value == u"是"):
            p = doc.tables[0].cell(12, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 19).value == u"否"):
            p = doc.tables[0].cell(12, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(12, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充上部结构第一项")
            p.font.name = u"Segoe UI Symbol"

        # 是否存在渗漏、裂缝情况
        p = doc.tables[0].cell(13, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 22).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 21).value == u"是"):
            p = doc.tables[0].cell(13, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 21).value == u"否"):
            p = doc.tables[0].cell(13, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(13, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充屋面第一项")

            p.font.name = u"Segoe UI Symbol"

        # 女儿墙高度、防护栏杆是否满足要求
        p = doc.tables[0].cell(14, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 24).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 23).value == u"是"):
            p = doc.tables[0].cell(14, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 23).value == u"否"):
            p = doc.tables[0].cell(14, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(14, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充屋面第二项")

            p.font.name = u"Segoe UI Symbol"

        # 是否存在加层或搭建物（如广告牌、钢棚等）而影响房屋使用安全情况
        p = doc.tables[0].cell(15, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 26).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 25).value == u"是"):
            p = doc.tables[0].cell(15, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 25).value == u"否"):
            p = doc.tables[0].cell(15, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(15, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充屋面第三项")

            p.font.name = u"Segoe UI Symbol"

        # 外墙及外墙上附属构件是否存在松脱、开裂、腐蚀等情况
        p = doc.tables[0].cell(16, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 28).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 27).value == u"是"):
            p = doc.tables[0].cell(16, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 27).value == u"否"):
            p = doc.tables[0].cell(16, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(16, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充外围结构")
            p.font.name = u"Segoe UI Symbol"

        # 天花吊顶是否存在变形、脱落情况
        p = doc.tables[0].cell(17, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 30).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 29).value == u"是"):
            p = doc.tables[0].cell(17, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 29).value == u"否"):
            p = doc.tables[0].cell(17, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(17, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充室内第一项")

            p.font.name = u"Segoe UI Symbol"

        # 临边等位置的防护栏杆设置和高度是否满足要求
        p = doc.tables[0].cell(18, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 32).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 31).value == u"是"):
            p = doc.tables[0].cell(18, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 31).value == u"否"):
            p = doc.tables[0].cell(18, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(18, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充室内第二项")

            p.font.name = u"Segoe UI Symbol"

        # 是否存在房屋主体结构受自然灾害以及爆炸、火灾等事故损坏情况
        p = doc.tables[0].cell(19, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 34).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 33).value == u"是"):
            p = doc.tables[0].cell(19, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 33).value == u"否"):
            p = doc.tables[0].cell(19, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(19, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充其他第一项")

            p.font.name = u"Segoe UI Symbol"

        # 是否有房屋拆改主体或者承重结构、改变使用功能或者明显加大荷载等情况
        p = doc.tables[0].cell(20, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 36).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 35).value == u"是"):
            p = doc.tables[0].cell(20, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 35).value == u"否"):
            p = doc.tables[0].cell(20, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(20, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充其他第二项")

            p.font.name = u"Segoe UI Symbol"

        # 建房是否有削坡或进行人工边坡，削坡或人工边坡的安全情况
        p = doc.tables[0].cell(21, 7).paragraphs[0].clear().add_run(
            self._all_to_string(excel_cell.cell(i, 38).value))  # text="aaaa"
        p.font.name = u"宋体"
        if (excel_cell.cell(i, 37).value == u"是"):
            p = doc.tables[0].cell(21, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 37).value == u"否"):
            p = doc.tables[0].cell(21, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(21, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            self._list_problem.append("补充其他第三项")

            p.font.name = u"Segoe UI Symbol"

        if (excel_cell.cell(i, 39).value == u"是"):
            p = doc.tables[0].cell(22, 5).paragraphs[0].clear().add_run("是☑  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        elif (excel_cell.cell(i, 39).value == u"否"):
            p = doc.tables[0].cell(22, 5).paragraphs[0].clear().add_run("是□  否☑")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        else:
            p = doc.tables[0].cell(22, 5).paragraphs[0].clear().add_run("是□  否□")  # text="aaaa"
            p.font.name = u"Segoe UI Symbol"
        if (excel_cell.cell(i, 39).value == u"是"):
            p = doc.paragraphs[31].add_run(self._all_to_string(excel_cell.cell(i, 9).value))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.font.name = u"宋体"
            p.font.size = Pt(14)
            p.font.bold = True

            p = doc.paragraphs[32].add_run(self._all_to_string(excel_cell.cell(i, 3).value))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.font.name = u"宋体"
            p.font.size = Pt(14)
            p.font.bold = True

            a = re.search("(.*?)[(（](.*?)[)）]", excel_cell.cell(i, 0).value)
            if a:
                p = doc.paragraphs[33].add_run(a.group(0));
                p.font.name = u"宋体"
                p.font.size = Pt(14)
                p.font.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                p = doc.paragraphs[33].add_run(self._all_to_string(excel_cell.cell(i, 0).value));
                p.font.name = u"宋体"
                p.font.size = Pt(14)
                p.font.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            p = doc.paragraphs[4].add_run(self._all_to_string(excel_cell.cell(i, 40).value))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.font.name = u"宋体"
            p.font.size = Pt(12)

            p = doc.paragraphs[5].add_run(self._all_to_string(excel_cell.cell(i, 41).value))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.font.name = u"宋体"
            p.font.size = Pt(12)
            # p.font.bold = True

            p = doc.paragraphs[6].add_run(self._all_to_string(excel_cell.cell(i, 42).value))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.font.name = u"宋体"
            p.font.size = Pt(12)

            p = doc.paragraphs[15].add_run(self._all_to_string(excel_cell.cell(i, 9).value))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.font.name = u"宋体"
            p.font.size = Pt(14)
            p.font.bold = True

            p = doc.paragraphs[16].add_run(self._all_to_string(excel_cell.cell(i, 3).value))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.font.name = u"宋体"
            p.font.size = Pt(14)
            p.font.bold = True


            a = re.search("(.*?)[(（](.*?)[)）]", excel_cell.cell(i, 0).value)
            if a:
                p = doc.paragraphs[17].add_run(a.group(1));
                p.font.name = u"宋体"
                p.font.size = Pt(14)
                p.font.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                p = doc.paragraphs[17].add_run(self._all_to_string(excel_cell.cell(i, 0).value));
                p.font.name = u"宋体"
                p.font.size = Pt(14)
                p.font.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # p.font.name = u"宋体"
    def _insert_pictrue(self,pictrue_folder, i, doc,excel_cell):

        a = re.search("(.*?)[(（](.*?)[)）]", excel_cell.cell(i, 0).value)
        if a:
            pictrue_position = os.path.join(pictrue_folder,a.group(1))
        else:
            pictrue_position =os.path.join(pictrue_folder,excel_cell.cell(i, 0).value)
        if os.path.exists(pictrue_position) and os.path.isdir(pictrue_position):

            pictrue_position = pictrue_position
            pictrue_set = os.listdir(pictrue_position)
            limian=[]
            other=[]
            for iy in pictrue_set:
                # print(iy)
                img_path=os.path.join(os.getcwd(),pictrue_position,iy)
                if re.search("建筑物立面",iy):
                    limian.append(iy)
                    f = Image.open(img_path)  # 你的图片文件
                    f.save(img_path)  # 替换掉你的图片文件
                    f.close()
                else:
                    other.append(iy)
                    f = Image.open(img_path)  # 你的图片文件
                    f.save(img_path)  # 替换掉你的图片文件
                    f.close()
            limian.extend(other)
            pictrue_set=[]
            pictrue_set=limian
            # print(limian,other)
            # print(i,type(i))
            if (excel_cell.cell(int(i), 39).value == u"是"):
                for ii in range(len(pictrue_set)):
                    if ii == 0:
                        paragraph = doc.tables[2].cell(0, 0).paragraphs[0]
                        a = paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        # print(a.height)
                        if a.height > 3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))
                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[2].cell(1, 0).paragraphs[0].clear().add_run(
                            re.sub(r"\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
                    if ii == 1:
                        paragraph = doc.tables[2].cell(0, 1).paragraphs[0]
                        a = paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        if a.height > 3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))
                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[2].cell(1, 1).paragraphs[0].clear().add_run(
                            re.sub(r"\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
                    if ii == 2:
                        paragraph = doc.tables[2].cell(2, 0).paragraphs[0]
                        a = paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        if a.height > 3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))
                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[2].cell(3, 0).paragraphs[0].clear().add_run(
                            re.sub("\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
                    if ii == 3:
                        paragraph = doc.tables[2].cell(2, 1).paragraphs[0]
                        a = paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        if a.height > 3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))
                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[2].cell(3, 1).paragraphs[0].clear().add_run(
                            re.sub("\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
            else:

                for ii in range(len(pictrue_set)):
                    if ii == 0:
                        paragraph = doc.tables[1].cell(0, 0).paragraphs[0]
                        a =paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        # print(a.height)
                        if a.height>3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))
                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[1].cell(1, 0).paragraphs[0].clear().add_run(
                            re.sub(r"\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
                    if ii == 1:
                        paragraph = doc.tables[1].cell(0, 1).paragraphs[0]
                        a=paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        if a.height>3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))
                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[1].cell(1, 1).paragraphs[0].clear().add_run(
                            re.sub(r"\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
                    if ii == 2:
                        paragraph = doc.tables[1].cell(2, 0).paragraphs[0]
                        a=paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        if a.height>3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))
                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[1].cell(3, 0).paragraphs[0].clear().add_run(
                            re.sub("\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
                    if ii == 3:
                        paragraph = doc.tables[1].cell(2, 1).paragraphs[0]
                        a=paragraph.clear().add_run().add_picture(
                            os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.2))
                        if a.height>3000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(1.7))

                        elif a.height<2000000:
                            paragraph.clear().add_run().add_picture(
                                os.path.join(pictrue_position, pictrue_set[ii]), width=Inches(2.6))
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p = doc.tables[1].cell(3, 1).paragraphs[0].clear().add_run(
                            re.sub("\d*", "", os.path.splitext(pictrue_set[ii])[0]))
                        p.font.name = u"宋体"
            self._pictrue_count=len(pictrue_set)
        else:
            # print(excel_cell.cell(i, 0).value, "不存在照片")
            self._list_problem.append("缺少照片")
    def check_filename_available(self,filename):
        n = [0]

        def check_meta(file_name):
            file_name_new = file_name
            if os.path.isfile(file_name):
                file_name_new = file_name[:file_name.rfind('.')] + '重复_' + str(n[0]) + file_name[file_name.rfind('.'):]
                n[0] += 1
            if os.path.isfile(file_name_new):
                file_name_new = check_meta(file_name)
            return file_name_new

        return_name = check_meta(filename)
        return return_name
    def _word_all(self,i):
        # print(i)
        if (self._excel_cell.cell(i, 39).value == u"是"):
            doc = docx.Document(self._docx_path2)
        else:
            doc = docx.Document(self._docx_path1)
        self._insert_word(self._excel_cell, i, doc)
        self._insert_pictrue(self._pictrue_folder, i, doc,self._excel_cell)
        # doc.save(os.path.join(self._docx_folder, self._all_to_string(self._excel_cell.cell(i, 0).value) + r'.docx'))
        self._save_p=os.path.join(self._docx_folder, self._all_to_string(self._excel_cell.cell(i, 0).value) + r'.docx')
        # print(self.check_filename_available(self._save_p))
        doc.save(self.check_filename_available(self._save_p))
        return 1
# he(
    def _main(self):

          # 通过名称获取
        start=time.time()
        for i in range(2, self._excel_cell.nrows):
            self._word_all(i)
        end=time.time()
        # print(end-start)
class write_excel():
    def __init__(self):
        self.workbook = xlwt.Workbook(encoding='utf-8')
        # 创建一个worksheet
        self.worksheet = self.workbook.add_sheet('My Worksheet',cell_overwrite_ok=True)
        self.worksheet.write(0, 1, "普查编号")
        self.worksheet.write(0, 2, "问题")
        self.worksheet.write(0, 3,  "照片数量")
        self.worksheet.write(0, 4,)
        # pass

    def execute_main(self):

        # workbook = xlwt.Workbook(encoding='utf-8')        print(os.path.join(self._s1,"excel_test.xls"))


        # print(len(w_sheet.rows))
        # for j in range(len(self.w_sheet.rows), self._excel_cell.nrows-2 + len(self.w_sheet.rows)):
        #     print("k",j)
        #     print(self._serial,self._list_problem)len(self.w_sheet.rows)len(self.w_sheet.rows)++
        # print(len(self.w_sheet.rows))
        if self._list_problem or self._pictrue_count:
            print(self._i,self._serial,",".join(self._list_problem),str(self._pictrue_count))
            self.worksheet.write(self._i-1, 1, self._i)
            self.worksheet.write(self._i-1, 1, self._serial)
            self.worksheet.write(self._i-1, 2, ",".join(self._list_problem))
            self.worksheet.write(self._i-1, 3, str(self._pictrue_count))
            # self.w_sheet.write(self._i-1, 4, str(self._pictrue_count))


class Reg(Frame,deal_word,write_excel):

    def __init__(self, master):
        self.success=0
        self.failue=0
        self.frame = Frame(master)
        self.frame.pack()
        self.lab1 = Label(self.frame, text="路径:")
        self.lab1.grid(row=0, column=0, sticky=E)
        self.ent1 = Entry(self.frame)
        self.ent1.grid(row=0, column=1, sticky=W)
        self.lab2 = Label(self.frame, text="待用:")
        self.lab2.grid(row=1, column=0,sticky=E)
        self.ent2 = Entry(self.frame, show="*")
        self.failure=[]
        self.ent2.grid(row=1, column=1, sticky=W)
        self.button = Button(self.frame, text="开始", command=self.Submit)
        self.button.grid(row=1, column=2, sticky=W,ipadx=20)
        self.lab3 = Label(self.frame, text="")
        self.lab3.grid(row=3, column=1, sticky=W)
        self.lab3.grid(row=0,column=2,sticky=W)
        self.mpb = Progressbar(self.frame, orient="horizontal", length=200, value=0, mode="determinate")
        self.mpb.grid(row=2, column=0, columnspan=2)

        deal_word.__init__(self)
        write_excel.__init__(self)


        self.scrolW = 40  # 设置文本框的长度
        self.scrolH = 18  # 设置文本框的高度
        self.text = scrolledtext.ScrolledText(self.frame, width=self.scrolW, height=self.scrolH, wrap=WORD)
        self.text.grid(row=3,columnspan=3)


    def Submit(self):
        self.text.insert("end","")
        s1 = self.ent1.get()
             # or r"C:\Users\道路\Desktop\模板four\凤岗内业\非危房"
        s2 = self.ent2.get()
        if os.path.exists(s1):
            os.chdir(s1)
            self._s1=s1
            self._excel_cell = xlrd.open_workbook(self._excel_path).sheet_by_name("数据源")

            # cl = deal_word(s1)
            # cl._main()
            start = time.time()
            self.mpb["maximum"] =self._excel_cell.nrows-1
            # try:
            for i in range(2, self._excel_cell.nrows):


                # try:
                    self._serial=self._excel_cell.cell(i, 0).value
                    self._word_all(i)
                    self.mpb["value"] = i
                    self.text.see(END)  # 一直查看文本的最后位置~
                    # print(all_to_string(sht[i, 0].value))
                    self.text.insert("end", str(i)+":"+self._all_to_string(self._excel_cell.cell(i, 0).value)+'成功生成Word'+ "\n"+"需要补充:"+str(self._list_problem)+"\n"+"照片数量:"+str(self._pictrue_count)+"\n\n")
                    root.update()
                    self.success+=1
                    if self._list_problem or self._pictrue_count:
                        self._i = i
                        self.execute_main()
                    self._list_problem=[]
                    self._pictrue_count=0
                    self._serial=""
                # except Exception as e:
                #     self._serial=self._excel_cell.cell(i, 0).value
                #     # print(e)
                #     self.failure.append([self._all_to_string(self._excel_cell.cell(i, 0).value),"错误原因:"+str(e)])
                #     self.failue+=1
                #     self._list_problem=[]
                #     self._pictrue_count=0
                #     self._serial=""
                #     pass


            self.text.insert('end',r"程序运行完成：总数"+str(self._excel_cell.nrows-2)+ "\n\n")
            self.text.insert("end","成功"+str(self.success)+ "\n\n")
            self.text.insert("end", "失败" +  str(self.failue)+"\r\n"+"失败原因:"+str(self.failure) + "\n\n")
            self.text.update()


            end=time.time()

            self.text.see(END)  # 一直查看文本的最后位置~
            root.update()
            self.text.insert('end',"运行时间"+str(end - start)+"s"+ "\r\n")  # 结束时间-开始时间
            self.workbook.save(os.path.join(self._s1,"excel_test.xls"))



        else:
            self.lab3["text"] = "请输入路径!"
        # self.ent1.delete(0, len(s1))
        # self.ent2.delete(0, len(s2))


root = Tk()
root.title("凤岗内业")
# root.iconbitmap("jpg.ico")

# root.geometry("500x500")
root.resizable(width=False, height=False)
app = Reg(root)






root.mainloop()
